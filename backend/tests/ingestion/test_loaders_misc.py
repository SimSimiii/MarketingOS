import pytest
from docx import Document as DocxDocument

from app.ingestion.documents import SourceType
from app.ingestion.exceptions import LoaderError
from app.ingestion.loaders.docx_loader import DocxLoader
from app.ingestion.loaders.json_loader import JsonLoader
from app.ingestion.loaders.markdown_loader import MarkdownLoader
from app.ingestion.loaders.text_loader import PlainTextLoader


@pytest.mark.asyncio
async def test_markdown_loader_accepts_raw_text():
    raw = await MarkdownLoader().load("# Title\n\nBody text.")
    assert raw.source_type == SourceType.MARKDOWN
    assert raw.content == "# Title\n\nBody text."


@pytest.mark.asyncio
async def test_markdown_loader_reads_from_file(tmp_path):
    path = tmp_path / "doc.md"
    path.write_text("# From file\n", encoding="utf-8")
    raw = await MarkdownLoader().load(str(path))
    assert raw.content == "# From file"


@pytest.mark.asyncio
async def test_plain_text_loader_accepts_raw_text():
    raw = await PlainTextLoader().load("just some text")
    assert raw.source_type == SourceType.PLAIN_TEXT
    assert raw.content == "just some text"


@pytest.mark.asyncio
async def test_json_loader_pretty_prints_raw_json():
    raw = await JsonLoader().load('{"a": 1, "b": [2, 3]}')
    assert raw.source_type == SourceType.JSON
    assert '"a": 1' in raw.content


@pytest.mark.asyncio
async def test_json_loader_reads_from_file(tmp_path):
    path = tmp_path / "data.json"
    path.write_text('{"key": "value"}', encoding="utf-8")
    raw = await JsonLoader().load(str(path))
    assert '"key": "value"' in raw.content


@pytest.mark.asyncio
async def test_json_loader_invalid_json_raises_loader_error():
    with pytest.raises(LoaderError):
        await JsonLoader().load("{not valid json")


@pytest.mark.asyncio
async def test_docx_loader_extracts_headings_and_paragraphs(tmp_path):
    path = tmp_path / "doc.docx"
    document = DocxDocument()
    document.add_heading("My Heading", level=1)
    document.add_paragraph("A regular paragraph.")
    document.save(str(path))

    raw = await DocxLoader().load(str(path))

    assert raw.source_type == SourceType.DOCX
    assert "# My Heading" in raw.content
    assert "A regular paragraph." in raw.content


@pytest.mark.asyncio
async def test_docx_loader_missing_file_raises_loader_error(tmp_path):
    with pytest.raises(LoaderError):
        await DocxLoader().load(str(tmp_path / "missing.docx"))
