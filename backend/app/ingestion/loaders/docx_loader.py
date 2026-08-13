from datetime import UTC, datetime
from pathlib import Path

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from app.ingestion.documents import RawDocument, SourceType
from app.ingestion.exceptions import LoaderError
from app.ingestion.loaders.base import Loader


class DocxLoader(Loader):
    """`source` is a path to a .docx file. Paragraphs styled as headings are
    rendered as Markdown headings so downstream heading-based chunking works;
    everything else is plain paragraph text."""

    source_type = SourceType.DOCX

    async def load(self, source: str) -> RawDocument:
        path = Path(source)
        if not path.is_file():
            raise LoaderError(f"DOCX source not found: '{source}'")

        try:
            document = DocxDocument(str(path))
        except PackageNotFoundError as exc:
            raise LoaderError(f"Failed to read DOCX '{source}': {exc}") from exc

        lines: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name or "") if paragraph.style else ""
            if style.startswith("Heading"):
                level = "".join(ch for ch in style if ch.isdigit()) or "1"
                lines.append(f"{'#' * int(level)} {text}")
            else:
                lines.append(text)

        core_props = document.core_properties
        return RawDocument(
            content="\n\n".join(lines),
            source=source,
            source_type=self.source_type,
            fetched_at=datetime.now(UTC),
            metadata={
                "title": (core_props.title or "").strip(),
                "author": (core_props.author or "").strip(),
            },
        )
